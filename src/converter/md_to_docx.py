"""Markdown → DOCX converter.

Handles:
- Headings (# to ######)
- Bold (**text**), Italic (*text*), Bold+Italic (***text***)
- Unordered lists (- / * / +) with nesting
- Ordered lists (1. 2. 3.) with nesting
- Tables (GFM pipe syntax)
- Blockquotes (>)
- Inline code (`code`) and fenced code blocks (``` ```)
- Horizontal rules (---, ***)
- Plain paragraphs
"""
from __future__ import annotations

import os
import re

from .base import BaseConverter, ConversionResult


class MdToDocxConverter(BaseConverter):
    FROM_EXT = ".md"
    TO_EXT = ".docx"

    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            import docx
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        # Strip YAML frontmatter
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                content = parts[2].lstrip("\n")

        document = docx.Document()
        self._apply_styles(document)

        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # Fenced code block
            if line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                para = document.add_paragraph("\n".join(code_lines), style="No Spacing")
                para.runs[0].font.name = "Courier New"
                para.runs[0].font.size = Pt(10)
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^(\s*[-*_]){3,}\s*$", line):
                document.add_paragraph("─" * 40, style="No Spacing")
                i += 1
                continue

            # Heading
            m = re.match(r"^(#{1,6})\s+(.+)$", line)
            if m:
                level = len(m.group(1))
                document.add_heading(m.group(2).strip(), level=level)
                i += 1
                continue

            # Table — collect all consecutive table lines
            if re.match(r"^\|.+\|", line):
                table_lines = []
                while i < len(lines) and re.match(r"^\|.+\|", lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(document, table_lines)
                continue

            # Blockquote
            if line.startswith(">"):
                text = re.sub(r"^>\s?", "", line)
                para = document.add_paragraph(style="No Spacing")
                self._add_inline(para, text)
                para.paragraph_format.left_indent = docx.shared.Inches(0.4)
                i += 1
                continue

            # Unordered list
            m = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
            if m:
                depth = len(m.group(1)) // 2
                style = "List Bullet" if depth == 0 else f"List Bullet {depth + 1}"
                try:
                    para = document.add_paragraph(style=style)
                except KeyError:
                    para = document.add_paragraph(style="List Bullet")
                self._add_inline(para, m.group(3))
                i += 1
                continue

            # Ordered list
            m = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
            if m:
                depth = len(m.group(1)) // 2
                style = "List Number" if depth == 0 else f"List Number {depth + 1}"
                try:
                    para = document.add_paragraph(style=style)
                except KeyError:
                    para = document.add_paragraph(style="List Number")
                self._add_inline(para, m.group(2))
                i += 1
                continue

            # Blank line
            if not line.strip():
                i += 1
                continue

            # Regular paragraph — collect continuation lines
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not self._is_block_start(lines[i]):
                para_lines.append(lines[i])
                i += 1
            para = document.add_paragraph()
            self._add_inline(para, " ".join(para_lines))

        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        document.save(output_path)

        return ConversionResult(
            output_path=output_path,
            success=True,
            message=f"Converted {os.path.basename(input_path)} → {os.path.basename(output_path)}",
        )

    # ── internal helpers ────────────────────────────────────────────────────

    def _is_block_start(self, line: str) -> bool:
        return bool(
            re.match(r"^#{1,6}\s", line)
            or re.match(r"^(\s*)([-*+])\s+", line)
            or re.match(r"^(\s*)\d+\.\s+", line)
            or line.startswith(">")
            or line.startswith("```")
            or re.match(r"^\|.+\|", line)
            or re.match(r"^(\s*[-*_]){3,}\s*$", line)
        )

    def _add_inline(self, para, text: str) -> None:
        """Add text with inline bold/italic formatting to a paragraph."""
        # Pattern: ***text***, **text**, *text*, `code`
        pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
        last = 0
        for m in pattern.finditer(text):
            # Plain text before match
            if m.start() > last:
                para.add_run(text[last:m.start()])
            full = m.group(0)
            if full.startswith("***"):
                run = para.add_run(m.group(2))
                run.bold = True
                run.italic = True
            elif full.startswith("**"):
                run = para.add_run(m.group(3))
                run.bold = True
            elif full.startswith("*"):
                run = para.add_run(m.group(4))
                run.italic = True
            elif full.startswith("`"):
                from docx.shared import Pt
                run = para.add_run(m.group(5))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            last = m.end()
        if last < len(text):
            para.add_run(text[last:])

    def _add_table(self, document, table_lines: list[str]) -> None:
        # Filter out separator rows (|---|---|)
        data_rows = [
            line for line in table_lines
            if not re.match(r"^\|[\s\-:|]+\|$", line)
        ]
        if not data_rows:
            return

        rows = []
        for line in data_rows:
            # Strip leading/trailing |
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(cells)

        col_count = max(len(r) for r in rows)
        table = document.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for ri, row_data in enumerate(rows):
            for ci, cell_text in enumerate(row_data):
                if ci < col_count:
                    cell = table.cell(ri, ci)
                    para = cell.paragraphs[0]
                    self._add_inline(para, cell_text)
                    if ri == 0:
                        for run in para.runs:
                            run.bold = True

    def _apply_styles(self, document) -> None:
        """Ensure built-in styles exist (Word creates them on demand)."""
        pass
