"""Plain text → DOCX converter.

Heuristics:
- First non-blank line → Heading 1 title
- Blank lines separate paragraphs
- Lines that look like section headers (all-caps, ends without period, short)
  are promoted to Heading 2
"""
from __future__ import annotations

import os
import re

from .base import BaseConverter, ConversionResult


class TxtToDocxConverter(BaseConverter):
    FROM_EXT = ".txt"
    TO_EXT = ".docx"

    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        document = docx.Document()

        lines = content.splitlines()
        title_used = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if not title_used and len(stripped) < 80 and not stripped.endswith("."):
                document.add_heading(stripped, level=1)
                title_used = True
            elif _looks_like_heading(stripped) and title_used:
                document.add_heading(stripped, level=2)
            else:
                document.add_paragraph(stripped)

        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        document.save(output_path)

        return ConversionResult(
            output_path=output_path,
            success=True,
            message=f"Converted {os.path.basename(input_path)} → {os.path.basename(output_path)}",
        )


def _looks_like_heading(line: str) -> bool:
    """Heuristic: short, no sentence-ending punctuation, not all lowercase."""
    if len(line) > 60:
        return False
    if line.endswith((".", "。", "!", "?", "！", "？")):
        return False
    # At least one uppercase or Chinese character, not all lowercase ascii
    if re.search(r"[A-Z\u4e00-\u9fff]", line):
        return True
    return False
