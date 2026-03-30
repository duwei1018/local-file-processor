"""PDF → DOCX converter.

Extracts text from each PDF page and creates a Word document with:
- Document title as Heading 1
- Each page's text as a normal paragraph block
"""
from __future__ import annotations

import os

from .base import BaseConverter, ConversionResult


class PdfToDocxConverter(BaseConverter):
    FROM_EXT = ".pdf"
    TO_EXT = ".docx"

    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            import pypdf
            import docx
        except ImportError:
            raise ImportError("pypdf and python-docx required: pip install pypdf python-docx")

        reader = pypdf.PdfReader(input_path)

        meta = reader.metadata or {}
        title = (
            meta.get("/Title")
            or meta.get("title")
            or os.path.splitext(os.path.basename(input_path))[0]
        )
        author = meta.get("/Author") or meta.get("author")

        document = docx.Document()
        document.add_heading(str(title), level=1)
        if author:
            document.add_paragraph(f"作者：{author}")

        for i, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            if len(reader.pages) > 1:
                document.add_heading(f"第 {i} 页", level=2)
            for para_text in text.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    document.add_paragraph(para_text)

        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        document.save(output_path)

        return ConversionResult(
            output_path=output_path,
            success=True,
            message=f"Converted {os.path.basename(input_path)} → {os.path.basename(output_path)}",
        )
