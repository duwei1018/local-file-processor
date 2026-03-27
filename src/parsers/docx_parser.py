import os
from .base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = docx.Document(file_path)

        paragraphs = []
        headings = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            if para.style.name.startswith("Heading"):
                headings.append(para.text.strip())
            paragraphs.append(para.text.strip())

        full_text = "\n\n".join(paragraphs)

        props = doc.core_properties
        title = props.title or os.path.splitext(os.path.basename(file_path))[0]
        author = props.author or None

        return ParsedDocument(
            title=title,
            text=full_text,
            file_path=file_path,
            file_type="docx",
            author=author,
            headings=headings,
        )
